#!/usr/bin/env python3
"""读取文档文件并输出纯文本。支持 .docx / .txt / .md。

用法:
  python3 read_doc.py input.docx
  python3 read_doc.py input.txt
  python3 read_doc.py input.md
  echo "直接传入文本" | python3 read_doc.py -
"""

from __future__ import annotations

import sys
from pathlib import Path


def read_docx(path: str) -> str:
    """读取 .docx 文件，提取全部段落文本。"""
    try:
        from docx import Document
    except ImportError:
        print(
            "[ERROR] 需要 python-docx: pip3 install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def read_text(path: str) -> str:
    """读取纯文本文件（.txt / .md 等）。"""
    return Path(path).read_text(encoding="utf-8")


def main():
    if len(sys.argv) < 2:
        print("用法: python3 read_doc.py <文件路径>", file=sys.stderr)
        sys.exit(1)

    source = sys.argv[1]

    # 从 stdin 读取
    if source == "-":
        print(sys.stdin.read())
        return

    path = Path(source)
    if not path.exists():
        print(f"[ERROR] 文件不存在: {source}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()

    if suffix == ".docx":
        text = read_docx(str(path))
    elif suffix in (".txt", ".md", ".text", ".markdown"):
        text = read_text(str(path))
    elif suffix == ".doc":
        print(
            "[ERROR] 不支持旧版 .doc 格式，请先转换为 .docx",
            file=sys.stderr,
        )
        sys.exit(1)
    else:
        # 尝试当纯文本读取
        try:
            text = read_text(str(path))
        except Exception:
            print(f"[ERROR] 不支持的文件格式: {suffix}", file=sys.stderr)
            sys.exit(1)

    print(text)


if __name__ == "__main__":
    main()
